from io import BytesIO

from app_logging import setup_logger
from config.config import OLLAMA_ENABLED, TEXT_MODEL
from docx import Document
from faker import Faker
from fastapi import APIRouter, Request, Response
from utils.helper import generate_random_name
from utils.ollama import generate_document_with_ollama
from utils.content_manager import ContentManager

logger = setup_logger(__name__)
router = APIRouter()
fake = Faker()
cm = ContentManager(default="index", extension="doc")


def return_doc_file(request: Request) -> Response:
    """Return a Word document with AI- or Faker-generated content."""
    cm.resolve(request)

    if cm.is_storing():
        if cached := cm.load():
            return Response(content=cached, media_type="application/octet-stream")

    doc = Document()
    doc.add_heading(fake.sentence(), 0)

    content = None
    if OLLAMA_ENABLED:
        try:
            prompt = f"Create a document titled '{cm.file_name}' with several paragraphs on the topic given in the title. Only return the contents of the document."
            logger.info(f"Ollama prompt: {prompt}")
            content = generate_document_with_ollama(prompt, TEXT_MODEL)
            logger.info("AI-generated content added.") if content else logger.warning("Ollama returned no content.")
        except Exception as e:
            logger.error(f"Ollama error: {e}")

    if not content:
        content = fake.paragraph(nb_sentences=3) + "\n\n" + fake.paragraph()
        logger.info("Faker-generated content added.")

    doc.add_paragraph(content)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)

    media_types = {
        ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ".doc": "application/msword",
        ".dotx": "application/vnd.openxmlformats-officedocument.wordprocessingml.template",
        ".dot": "application/msword",
        ".docm": "application/vnd.ms-word.document.macroEnabled.12",
        ".dotm": "application/vnd.ms-word.template.macroEnabled.12",
        ".odt": "application/vnd.oasis.opendocument.text",
    }
    ext = "." + cm.file_name.split(".")[-1].lower()
    media_type = media_types.get(ext, "application/octet-stream")

    logger.info(f"Returning {cm.file_name} with media type: {media_type}")

    response = Response(content=buf.getvalue(), media_type=media_type)
    response.headers["Content-Disposition"] = f"inline; filename={cm.file_name}"

    if cm.is_storing():
        cm.save(buf.getvalue())

    return response


# register routes after function is defined
ROUTES = ["/doc", "/docx", "/docs", "/documents"]
for route in ROUTES:
    router.add_api_route(f"{route}", return_doc_file, methods=["GET", "POST"], tags=["Documents"])
    router.add_api_route(f"{route}/{{file_name:path}}", return_doc_file, methods=["GET", "POST"], tags=["Documents"])
